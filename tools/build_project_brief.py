from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "docs" / "artifacts"
OUT_DIR.mkdir(parents=True, exist_ok=True)

doc = Document()
section = doc.sections[0]
section.top_margin = Inches(0.75)
section.bottom_margin = Inches(0.75)
section.left_margin = Inches(0.85)
section.right_margin = Inches(0.85)

styles = doc.styles
styles["Normal"].font.name = "Arial"
styles["Normal"].font.size = Pt(10.5)

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.LEFT
run = title.add_run("FutureLight Real Product Implementation Brief")
run.font.name = "Arial"
run.font.size = Pt(22)
run.font.bold = True
run.font.color.rgb = RGBColor(30, 41, 59)

subtitle = doc.add_paragraph()
subtitle.add_run("React frontend, Rust backend, PostgreSQL, asset QA, and content readiness workflow.").italic = True

doc.add_heading("Current Build", level=1)
for item in [
    "React/Vite frontend renders the first learning dashboard and calls the Rust API.",
    "Rust/Axum backend exposes health, DB health, page summary, and home summary endpoints.",
    "PostgreSQL runs through Docker Compose on localhost:37432.",
    "imagegen produced the first course cover and it is used in the frontend.",
    "Custom Codex plugins check asset and content readiness.",
]:
    doc.add_paragraph(item, style="List Bullet")

doc.add_heading("Verification Gates", level=1)
table = doc.add_table(rows=1, cols=3)
table.style = "Table Grid"
hdr = table.rows[0].cells
hdr[0].text = "Gate"
hdr[1].text = "Tool"
hdr[2].text = "Status"
rows = [
    ("Frontend build", "npm run build", "Passing"),
    ("Frontend lint", "npm run lint", "Passing"),
    ("Backend compile", "cargo check", "Passing"),
    ("Database health", "Docker + /api/health/db", "Passing"),
    ("Asset QA", "futurelight-assets", "Passing"),
    ("Content QA", "futurelight-content-checker", "Passing"),
]
for row in rows:
    cells = table.add_row().cells
    for index, value in enumerate(row):
        cells[index].text = value

doc.add_heading("Next Implementation Slice", level=1)
for item in [
    "Replace temporary fixed API responses with PostgreSQL-backed tables and seed data.",
    "Implement React route structure for the twelve documented pages.",
    "Add sound preference state and persistent user settings.",
    "Expand course content importer from spreadsheet inventory.",
]:
    doc.add_paragraph(item, style="List Number")

doc.save(OUT_DIR / "futurelight-mvp-brief.docx")
